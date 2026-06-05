#!/usr/bin/env python3
"""Build a Word report for the 20BN-Jester MAX78000 experiments."""

from __future__ import annotations

import json
import math
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("outputs/manual-20260530-jester/word-report")
DATA_PATH = ROOT / "jester_training_report_data_20260530.json"
CHART_DIR = ROOT / "charts"
DOCX_PATH = ROOT / "max78000_jester_training_report.docx"
BOARD_MEASUREMENTS_PATH = Path("outputs/manual-20260530-jester/board-measurements/board_inference_measurements.json")
CONFUSION_DIR = Path("outputs/manual-20260530-jester/confusion-all")


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"
TEXT = "1F2937"

SYNTHESIS_FAILED_ARCHES = {
    "jestermotion12resxconv": "synth fail",
    "jestermotion12resfc192": "synth fail",
}

EXTERNAL_REFERENCE_MODELS = [
    {
        "id": "saetlan_20bn_jester_3d_resnet101",
        "display": "saetlan/20BN-jester 3D ResNet101",
        "arch": "3d_resnet101",
        "params": 85_443_099,
        "top1": 85.99,
        "category": "External reference, not MAX78000 deployable",
        "notes": (
            "Top-1 is reported by the saetlan/20BN-jester README. "
            "Parameter count is calculated from the repository's 3D ResNet101 bottleneck model "
            "for 16 frames, 64x96 RGB input, and 27 classes."
        ),
    },
    {
        "id": "lin_tsm_resnet50_8f_jester",
        "display": "Lin et al. TSM ResNet-50 8F (Jester)",
        "arch": "tsm_resnet50_8f",
        "params": 24_300_000,
        "top1": 97.0,
        "category": "External reference, not MAX78000 deployable",
        "notes": (
            "Lin, Gan, and Han report 97.0% Top-1 on Jester for TSM. "
            "TSM adds no parameters over the 2D backbone; the paper reports 24.3M parameters "
            "for the 8-frame TSM ResNet-50 configuration."
        ),
    },
]

STM32U5_DEPLOYABLE_POINTS = [
    {"label": "STM32 U5 80K", "params": 80_000, "top1": 61.4},
    {"label": "STM32 U5 273K", "params": 273_000, "top1": 72.9},
    {"label": "STM32 U5 722K", "params": 722_000, "top1": 84.38},
]

LOSS_PLOT_EPOCH_LIMITS = {
    "cached_tinygesture_full_finetune300_20260519_201404": 80,
}

COMBINED_RAW8_LOSS_IDS = (
    "cached_tinygesture_full_finetune300_20260519_201404",
    "ai8x_jesternet_qat_20260520_005821",
)


def safe_id(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9._-]+", "_", text)
    return text.strip("_")[:90]


def fmt_pct(value: float | None, digits: int = 3) -> str:
    if value is None:
        return "n/a"
    return f"{value:.{digits}f}%"


def fmt_num(value: float | int | None, digits: int = 2) -> str:
    if value is None:
        return "n/a"
    if isinstance(value, int):
        return f"{value:,}"
    return f"{value:.{digits}f}"


def fmt_time(seconds: float | None) -> str:
    if not seconds:
        return "not logged"
    hours = seconds / 3600.0
    if hours >= 24:
        return f"{hours / 24.0:.1f} d"
    if hours >= 1:
        return f"{hours:.1f} h"
    return f"{seconds / 60.0:.1f} min"


def load_board_measurements() -> tuple[dict[str, dict[str, Any]], list[dict[str, Any]]]:
    if not BOARD_MEASUREMENTS_PATH.is_file():
        return {}, []
    payload = json.loads(BOARD_MEASUREMENTS_PATH.read_text(encoding="utf-8-sig"))
    by_model: dict[str, dict[str, Any]] = {}
    measurements = payload.get("measurements") or []
    for record in measurements:
        for model_id in record.get("applies_to") or [record.get("model_id")]:
            if model_id:
                by_model[model_id] = record
    return by_model, measurements


def fmt_inference_ms(model: dict[str, Any], board_by_model: dict[str, dict[str, Any]]) -> str:
    if not model.get("deployable") or model.get("mmac") is None:
        return "n/a"
    record = board_by_model.get(model["id"])
    if record and record.get("median_ms") is not None:
        return f"{float(record['median_ms']):.3f}"
    return SYNTHESIS_FAILED_ARCHES.get(model.get("arch"), "not measured")


def load_confusion_summary() -> dict[str, dict[str, str]]:
    path = CONFUSION_DIR / "summary.csv"
    if not path.is_file():
        return {}
    import csv

    rows: dict[str, dict[str, str]] = {}
    with path.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows[row["model"]] = row
    return rows


def phase_epoch_count(phase: dict[str, Any]) -> int:
    return len(phase.get("epochs") or [])


def qat_phases(model: dict[str, Any]) -> list[dict[str, Any]]:
    return [p for p in model.get("phases", []) if p.get("name", "").startswith("qat")]


def best_qat(model: dict[str, Any]) -> dict[str, Any] | None:
    candidates = []
    for phase in qat_phases(model):
        best = phase.get("best") or {}
        if best.get("top1") is not None:
            candidates.append((best["top1"], best.get("top5"), phase.get("name"), best.get("epoch")))
    if not candidates:
        return None
    top1, top5, phase, epoch = max(candidates, key=lambda item: item[0])
    return {"top1": top1, "top5": top5, "phase": phase, "epoch": epoch}


def plot_accuracy_value(model: dict[str, Any]) -> float | None:
    q = best_qat(model)
    if q and q.get("top1") is not None:
        return float(q["top1"])
    best = model.get("best") or {}
    if best.get("top1") is not None:
        return float(best["top1"])
    return None


def model_features(model: dict[str, Any]) -> dict[str, Any]:
    text = f"{model.get('id', '')} {model.get('display', '')} {model.get('arch', '')}".lower()
    arch = str(model.get("arch", "")).lower()
    is_teacher = "teacher" in arch or model.get("deployable") is False
    uses_motion = "motion" in text or "mixed" in text
    uses_signed = "signed" in text
    is_distilled = (not is_teacher) and ("distill" in text or "teacher student" in text or "teacher_qat" in text)
    if is_teacher:
        family = "teacher/software"
        marker = "diamond"
    elif uses_motion:
        family = "motion"
        marker = "square"
    else:
        family = "raw frames"
        marker = "circle"
    tags = ""
    if uses_motion:
        tags += "M"
    if uses_signed:
        tags += "S"
    if is_distilled:
        tags += "D"
    return {
        "family": family,
        "marker": marker,
        "uses_motion": uses_motion,
        "uses_signed": uses_signed,
        "is_distilled": is_distilled,
        "tags": tags,
    }


def deployment_style(model: dict[str, Any], board_by_model: dict[str, dict[str, Any]]) -> dict[str, str]:
    if model.get("external_reference"):
        return {"status": "external reference", "color": "#DC2626"}
    if not model.get("deployable"):
        return {"status": "software-only", "color": "#6B7280"}
    if SYNTHESIS_FAILED_ARCHES.get(model.get("arch")):
        return {"status": "synth failed", "color": "#F97316"}
    if board_by_model.get(model.get("id")):
        return {"status": "board/synth OK", "color": "#059669"}
    return {"status": "ai8x, not board-timed", "color": "#2563EB"}


def deployment_draw_order(status: str) -> int:
    # Draw successful board/synth points last so they remain visible in overlapping clusters.
    return {
        "external reference": 0,
        "software-only": 1,
        "synth failed": 2,
        "ai8x, not board-timed": 3,
        "board/synth OK": 4,
        "STM32 U5 deployable": 5,
    }.get(status, 1)


def draw_marker(draw: ImageDraw.ImageDraw, x: float, y: float, color: str, marker: str, radius: int = 8) -> None:
    if marker == "square":
        draw.rectangle((x - radius, y - radius, x + radius, y + radius), fill=color, outline="#111827")
    elif marker == "triangle":
        draw.polygon([(x, y - radius - 2), (x - radius - 2, y + radius), (x + radius + 2, y + radius)], fill=color, outline="#111827")
    elif marker == "diamond":
        draw.polygon([(x, y - radius - 3), (x - radius - 3, y), (x, y + radius + 3), (x + radius + 3, y)], fill=color, outline="#111827")
    else:
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=color, outline="#111827")


def draw_feature_tags(draw: ImageDraw.ImageDraw, x: float, y: float, tags: str, font: ImageFont.ImageFont) -> None:
    if not tags:
        return
    draw.text((x + 10, y + 5), tags, fill="#111827", font=font)


def fmt_params_short(value: int | float) -> str:
    value = float(value)
    if value >= 1_000_000:
        return f"{value / 1_000_000:.1f}M"
    return f"{value / 1_000:.0f}k"


def model_sort_key(model: dict[str, Any]) -> tuple[int, float]:
    bq = best_qat(model)
    top = bq["top1"] if bq else (model.get("best", {}) or {}).get("top1", -1)
    return (0 if model.get("deployable") else 1, -(top or -1))


def get_font(size: int = 18) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    ]
    for path in candidates:
        if Path(path).is_file():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_line(draw: ImageDraw.ImageDraw, points: list[tuple[float, float]], fill: str, width: int = 3) -> None:
    if len(points) < 2:
        return
    draw.line(points, fill=fill, width=width, joint="curve")
    r = 3
    for x, y in points:
        draw.ellipse((x - r, y - r, x + r, y + r), fill=fill)


def plot_loss_curve(model: dict[str, Any], out_path: Path) -> None:
    width, height = 1200, 610
    margin_l, margin_r, margin_t, margin_b = 92, 40, 70, 92
    chart_w = width - margin_l - margin_r
    chart_h = height - margin_t - margin_b
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = get_font(19)
    small = get_font(16)
    tiny = get_font(13)

    series: list[tuple[str, list[float], list[float], list[int]]] = []
    phase_ranges: list[tuple[str, int, int]] = []
    x_offset = 0
    all_values: list[float] = []
    epoch_limit = LOSS_PLOT_EPOCH_LIMITS.get(model.get("id", ""))
    for phase in model.get("phases", []):
        rows = phase.get("epochs") or []
        if epoch_limit is not None:
            remaining = epoch_limit - x_offset
            if remaining <= 0:
                break
            rows = rows[:remaining]
        if not rows:
            continue
        xs = list(range(x_offset + 1, x_offset + len(rows) + 1))
        train = [row.get("train_loss") for row in rows]
        val = [row.get("val_loss") for row in rows]
        train_pairs = [(x, y) for x, y in zip(xs, train) if y is not None]
        val_pairs = [(x, y) for x, y in zip(xs, val) if y is not None]
        if train_pairs:
            series.append((f"{phase['name']} train", [x for x, _ in train_pairs], [float(y) for _, y in train_pairs], xs))
            all_values.extend(float(y) for _, y in train_pairs)
        if val_pairs:
            series.append((f"{phase['name']} val", [x for x, _ in val_pairs], [float(y) for _, y in val_pairs], xs))
            all_values.extend(float(y) for _, y in val_pairs)
        phase_ranges.append((phase["name"], x_offset + 1, x_offset + len(rows)))
        x_offset += len(rows)

    if not series or not all_values:
        draw.text((margin_l, margin_t), "No epoch-level loss data was logged for this run.", fill="#555555", font=font)
        img.save(out_path)
        return

    x_min, x_max = 1, max(x_offset, 2)
    y_min = max(0.0, min(all_values) * 0.92)
    y_max = max(all_values) * 1.08
    if abs(y_max - y_min) < 1e-6:
        y_max = y_min + 1.0

    def tx(x: float) -> float:
        return margin_l + (x - x_min) / (x_max - x_min) * chart_w

    def ty(y: float) -> float:
        return margin_t + (y_max - y) / (y_max - y_min) * chart_h

    title = model["display"]
    draw.text((margin_l, 24), title[:95], fill=f"#{TEXT}", font=font)
    if epoch_limit is not None:
        draw.text((width - margin_r - 155, 27), f"first {epoch_limit} epochs shown", fill="#4B5563", font=tiny)

    # grid and axes
    draw.rectangle((margin_l, margin_t, margin_l + chart_w, margin_t + chart_h), outline=f"#{BORDER}", width=2)
    for i in range(6):
        y = margin_t + chart_h * i / 5
        value = y_max - (y_max - y_min) * i / 5
        draw.line((margin_l, y, margin_l + chart_w, y), fill="#EEF2F6", width=1)
        draw.text((12, y - 9), f"{value:.2f}", fill="#555555", font=tiny)
    draw.text((margin_l + chart_w / 2 - 30, height - 42), "Epoch", fill="#555555", font=small)
    draw.text((14, margin_t - 30), "Loss", fill="#555555", font=small)

    # phase boundaries
    for name, start, end in phase_ranges:
        if start > 1:
            x = tx(start - 0.5)
            draw.line((x, margin_t, x, margin_t + chart_h), fill="#C8D3DF", width=2)
        label_x = tx((start + end) / 2)
        draw.text((label_x - 30, margin_t + chart_h + 14), name[:18], fill="#666666", font=tiny)

    colors = ["#2563EB", "#F97316", "#059669", "#DC2626", "#7C3AED", "#0891B2", "#A16207", "#DB2777"]
    legend_x, legend_y = margin_l + 8, margin_t + 8
    for idx, (name, xs, ys, _) in enumerate(series):
        color = colors[idx % len(colors)]
        points = [(tx(x), ty(y)) for x, y in zip(xs, ys)]
        draw_line(draw, points, color, width=3)
        if idx < 8:
            y = legend_y + idx * 20
            draw.line((legend_x, y + 8, legend_x + 26, y + 8), fill=color, width=4)
            draw.text((legend_x + 34, y), name[:34], fill="#333333", font=tiny)

    img.save(out_path)


def build_raw8_combined_loss_model(model_by_id: dict[str, dict[str, Any]]) -> dict[str, Any] | None:
    finetune = model_by_id.get(COMBINED_RAW8_LOSS_IDS[0])
    qat = model_by_id.get(COMBINED_RAW8_LOSS_IDS[1])
    if not finetune or not qat:
        return None

    phases: list[dict[str, Any]] = []
    limit = LOSS_PLOT_EPOCH_LIMITS[COMBINED_RAW8_LOSS_IDS[0]]
    for phase in finetune.get("phases", []):
        rows = (phase.get("epochs") or [])[:limit]
        if rows:
            phases.append({**phase, "name": f"finetune 1-{limit}", "epochs": rows})
            break
    for phase in qat.get("phases", []):
        rows = phase.get("epochs") or []
        if rows:
            phases.append({**phase, "name": "ai8x QAT", "epochs": rows})

    if not phases:
        return None
    return {
        "id": "combined_tinygesture_finetune_jesternet_qat",
        "display": "TinyGesture raw 8-frame finetune + JesterNet ai8x QAT",
        "phases": phases,
    }


def add_training_curve_figure(
    doc: Document,
    chart_path: str,
    caption: str,
    figure_no: int,
    desc: str | None = None,
) -> None:
    if desc:
        doc.add_paragraph(desc)
    doc.add_picture(chart_path, width=Inches(6.25))
    cap = doc.add_paragraph(f"Figure {figure_no}. {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def plot_accuracy_mmac(models: list[dict[str, Any]], out_path: Path, board_by_model: dict[str, dict[str, Any]]) -> None:
    deployable = [m for m in models if m.get("deployable") and best_qat(m) and m.get("mmac")]
    width, height = 1300, 720
    margin_l, margin_r, margin_t, margin_b = 100, 62, 112, 96
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = get_font(21)
    small = get_font(15)
    tiny = get_font(12)
    draw.text((margin_l, 26), "Quantized Top-1 vs. Inference MMAC (ai8x models)", fill=f"#{TEXT}", font=font)
    draw.text(
        (margin_l, 54),
        "Color = deployment status; shape = input/model category; tags: M motion, S signed, D distilled.",
        fill="#4B5563",
        font=tiny,
    )
    if not deployable:
        img.save(out_path)
        return
    xs = [float(m["mmac"]) for m in deployable]
    ys = [float(best_qat(m)["top1"]) for m in deployable]
    x_min, x_max = min(xs) * 0.92, max(xs) * 1.08
    y_min, y_max = min(ys) - 1.5, max(ys) + 1.5
    chart_w, chart_h = width - margin_l - margin_r, height - margin_t - margin_b

    def tx(x: float) -> float:
        return margin_l + (x - x_min) / (x_max - x_min) * chart_w

    def ty(y: float) -> float:
        return margin_t + (y_max - y) / (y_max - y_min) * chart_h

    draw.rectangle((margin_l, margin_t, margin_l + chart_w, margin_t + chart_h), outline=f"#{BORDER}", width=2)
    for i in range(6):
        x = margin_l + chart_w * i / 5
        val = x_min + (x_max - x_min) * i / 5
        draw.line((x, margin_t, x, margin_t + chart_h), fill="#EEF2F6", width=1)
        draw.text((x - 18, margin_t + chart_h + 12), f"{val:.1f}", fill="#555555", font=tiny)
        y = margin_t + chart_h * i / 5
        yv = y_max - (y_max - y_min) * i / 5
        draw.line((margin_l, y, margin_l + chart_w, y), fill="#EEF2F6", width=1)
        draw.text((24, y - 8), f"{yv:.1f}%", fill="#555555", font=tiny)
    draw.text((margin_l + chart_w / 2 - 40, height - 42), "MMAC", fill="#555555", font=small)
    draw.text((18, margin_t - 28), "Top-1", fill="#555555", font=small)

    status_legend = [
        ("board/synth OK", "#059669"),
        ("ai8x, not board-timed", "#2563EB"),
        ("synth failed", "#F97316"),
    ]
    family_legend = [
        ("raw frames", "circle"),
        ("motion", "square"),
    ]
    legend_x, legend_y = margin_l + 10, margin_t + 10
    for idx, (label, color) in enumerate(status_legend):
        y = legend_y + idx * 22
        draw_marker(draw, legend_x, y + 8, color, "circle", radius=6)
        draw.text((legend_x + 17, y), label, fill="#333333", font=tiny)
    fam_x = legend_x + 178
    for idx, (label, marker) in enumerate(family_legend):
        y = legend_y + idx * 22
        draw_marker(draw, fam_x, y + 8, "#D1D5DB", marker, radius=6)
        draw.text((fam_x + 17, y), label, fill="#333333", font=tiny)
    draw.text((fam_x + 130, legend_y), "Tags: M=motion, S=signed, D=distilled", fill="#333333", font=tiny)

    draw_models = sorted(
        deployable,
        key=lambda m: (
            deployment_draw_order(deployment_style(m, board_by_model)["status"]),
            float(m.get("mmac") or 0),
            float((best_qat(m) or {}).get("top1") or 0),
        ),
    )
    for m in draw_models:
        q = best_qat(m)
        x, y = tx(float(m["mmac"])), ty(float(q["top1"]))
        features = model_features(m)
        style = deployment_style(m, board_by_model)
        draw_marker(draw, x, y, style["color"], features["marker"], radius=7)
        draw_feature_tags(draw, x, y, features["tags"], tiny)
        if m["display"] == "JesterNet raw 8-frame ai8x QAT":
            draw.text((x + 9, y - 11), "jesternet", fill="#374151", font=tiny)
    img.save(out_path)


def plot_accuracy_params(
    models: list[dict[str, Any]],
    out_path: Path,
    board_by_model: dict[str, dict[str, Any]],
    extra_points: list[dict[str, Any]] | None = None,
) -> None:
    points: list[dict[str, Any]] = []
    for model in models:
        top1 = plot_accuracy_value(model)
        params = model.get("params")
        if top1 is None or not params:
            continue
        features = model_features(model)
        style = deployment_style(model, board_by_model)
        points.append(
            {
                "label": model["display"],
                "arch": model.get("arch", ""),
                "params": int(params),
                "top1": top1,
                "status": style["status"],
                "color": style["color"],
                "family": features["family"],
                "marker": features["marker"],
                "tags": features["tags"],
            }
        )

    for ref in EXTERNAL_REFERENCE_MODELS:
        ref_model = {**ref, "deployable": False, "external_reference": True}
        features = model_features(ref_model)
        style = deployment_style(ref_model, board_by_model)
        points.append(
            {
                "label": ref["display"],
                "arch": ref["arch"],
                "params": int(ref["params"]),
                "top1": float(ref["top1"]),
                "status": style["status"],
                "color": style["color"],
                "family": features["family"],
                "marker": "diamond",
                "tags": "",
            }
        )

    for point in extra_points or []:
        points.append(
            {
                "label": point["label"],
                "arch": "stm32u5_reference",
                "params": int(point["params"]),
                "top1": float(point["top1"]),
                "status": "STM32 U5 deployable",
                "color": "#7C3AED",
                "family": "STM32 U5",
                "marker": "triangle",
                "tags": "",
            }
        )

    width, height = 1300, 740
    margin_l, margin_r, margin_t, margin_b = 112, 60, 112, 108
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = get_font(21)
    small = get_font(16)
    tiny = get_font(12)
    draw.text((margin_l, 28), "Top-1 Accuracy vs. Model Parameters", fill=f"#{TEXT}", font=font)
    draw.text(
        (margin_l, 54),
        "Color = deployment status/platform; shape = input/model category; tags: M motion, S signed, D distilled.",
        fill="#4B5563",
        font=tiny,
    )

    if not points:
        img.save(out_path)
        return

    x_values = [math.log10(p["params"]) for p in points]
    y_values = [p["top1"] for p in points]
    x_min, x_max = min(x_values) - 0.08, max(x_values) + 0.08
    y_min, y_max = max(0, min(y_values) - 4), min(100, max(y_values) + 4)
    chart_w, chart_h = width - margin_l - margin_r, height - margin_t - margin_b

    def tx(params: int) -> float:
        return margin_l + (math.log10(params) - x_min) / (x_max - x_min) * chart_w

    def ty(top1: float) -> float:
        return margin_t + (y_max - top1) / (y_max - y_min) * chart_h

    draw.rectangle((margin_l, margin_t, margin_l + chart_w, margin_t + chart_h), outline=f"#{BORDER}", width=2)

    x_ticks = [80_000, 200_000, 300_000, 500_000, 1_000_000, 3_000_000, 10_000_000, 30_000_000, 100_000_000]
    for tick in x_ticks:
        lx = math.log10(tick)
        if lx < x_min or lx > x_max:
            continue
        x = tx(tick)
        draw.line((x, margin_t, x, margin_t + chart_h), fill="#EEF2F6", width=1)
        draw.text((x - 20, margin_t + chart_h + 13), fmt_params_short(tick), fill="#555555", font=tiny)

    y_step = 5
    start_y = math.floor(y_min / y_step) * y_step
    end_y = math.ceil(y_max / y_step) * y_step
    y_tick = start_y
    while y_tick <= end_y:
        if y_min <= y_tick <= y_max:
            y = ty(y_tick)
            draw.line((margin_l, y, margin_l + chart_w, y), fill="#EEF2F6", width=1)
            draw.text((36, y - 8), f"{y_tick:.0f}%", fill="#555555", font=tiny)
        y_tick += y_step

    draw.text((margin_l + chart_w / 2 - 92, height - 44), "Parameters (log scale)", fill="#555555", font=small)
    draw.text((30, margin_t - 28), "Top-1 accuracy", fill="#555555", font=small)

    label_allowlist = {
        "saetlan/20BN-jester 3D ResNet101",
        "Lin et al. TSM ResNet-50 8F (Jester)",
        "Big software teacher",
        "Huge software teacher",
        "STM32 U5 80K",
        "STM32 U5 273K",
        "STM32 U5 722K",
    }
    sorted_points = sorted(points, key=lambda p: (deployment_draw_order(p["status"]), p["params"], p["top1"]))
    for point in sorted_points:
        x, y = tx(point["params"]), ty(point["top1"])
        radius = 10 if point["marker"] == "diamond" else 7
        draw_marker(draw, x, y, point["color"], point["marker"], radius=radius)
        draw_feature_tags(draw, x, y, point["tags"], tiny)
        if point["label"] in label_allowlist:
            label = point["label"]
            if label.startswith("saetlan"):
                label = "saetlan 3D ResNet101"
            elif label.startswith("Lin et al. TSM"):
                label = "TSM ResNet-50 8F"
            elif label.startswith("STM32 U5"):
                label = label.replace("STM32 U5 ", "U5 ")
            label_x = x + 10
            label_y = y - 9
            if point["label"] == "Lin et al. TSM ResNet-50 8F (Jester)":
                label_x = x - 54
                label_y = y + 14
            elif point["status"] == "STM32 U5 deployable":
                label_x = x + 12
                label_y = y - 18
            if x > margin_l + chart_w - 180:
                label_x = x - 174
            draw.text((label_x, label_y), label[:28], fill="#374151", font=tiny)

    legend = [
        ("board/synth OK", "#059669", "circle"),
        ("ai8x, not board-timed", "#2563EB", "circle"),
        ("synth failed", "#F97316", "circle"),
        ("software-only", "#6B7280", "circle"),
        ("external reference", "#DC2626", "circle"),
    ]
    if extra_points:
        legend.append(("STM32 U5 deployable", "#7C3AED", "triangle"))
    legend_x = margin_l + 10
    legend_y = margin_t + 12
    for idx, (label, color, marker) in enumerate(legend):
        y = legend_y + idx * 24
        draw_marker(draw, legend_x, y + 8, color, marker, radius=6)
        draw.text((legend_x + 18, y), label, fill="#333333", font=tiny)

    family_legend = [
        ("raw frames", "circle"),
        ("motion", "square"),
        ("teacher/external", "diamond"),
    ]
    family_x = legend_x + 190
    for idx, (label, marker) in enumerate(family_legend):
        y = legend_y + idx * 24
        draw_marker(draw, family_x, y + 8, "#D1D5DB", marker, radius=6)
        draw.text((family_x + 18, y), label, fill="#333333", font=tiny)
    draw.text((family_x + 160, legend_y), "Tags: M=motion, S=signed, D=distilled", fill="#333333", font=tiny)

    img.save(out_path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.0, color: str = "000000") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size: float = 7.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].width = Inches(widths[idx])
        set_cell_shading(header_cells[idx], LIGHT_GRAY)
        set_cell_text(header_cells[idx], header, bold=True, size=font_size, color="1F2937")
    for row_values in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for idx, value in enumerate(row_values):
            row.cells[idx].width = Inches(widths[idx])
            set_cell_text(row.cells[idx], value, size=font_size)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("MAX78000 20BN-Jester Training and Quantization Report")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Training curves, quantization results, MMAC, measured board timing, and confusion matrices")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def make_accuracy_rows(models: list[dict[str, Any]], board_by_model: dict[str, dict[str, Any]]) -> list[list[str]]:
    rows = []
    for m in sorted(models, key=model_sort_key):
        q = best_qat(m)
        b = m.get("best") or {}
        rows.append(
            [
                m["display"],
                "yes" if m.get("deployable") else "no",
                fmt_num(m.get("params"), 0),
                fmt_num(m.get("mmac"), 2),
                fmt_inference_ms(m, board_by_model),
                fmt_pct(b.get("top1")),
                fmt_pct(b.get("top5")),
                fmt_pct(q.get("top1") if q else None),
                fmt_pct(q.get("top5") if q else None),
            ]
        )
    return rows


def make_time_rows(models: list[dict[str, Any]]) -> list[list[str]]:
    rows = []
    for m in models:
        phases = m.get("phases") or []
        phase_bits = []
        for phase in phases:
            phase_bits.append(f"{phase['name']}: {phase_epoch_count(phase)} ep, {fmt_time(phase.get('duration_s'))}")
        rows.append([m["display"], "; ".join(phase_bits), fmt_time(m.get("training_time_s"))])
    return rows


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    models = data["models"]
    model_by_id = {model["id"]: model for model in models}
    board_by_model, board_measurements = load_board_measurements()
    confusion_summary = load_confusion_summary()

    for model in models:
        chart_path = CHART_DIR / f"{safe_id(model['id'])}.png"
        plot_loss_curve(model, chart_path)
        model["chart_path"] = str(chart_path)
        confusion_path = CONFUSION_DIR / f"{safe_id(model['id'])}_confusion_normalized.png"
        if confusion_path.is_file():
            model["confusion_path"] = str(confusion_path)
    combined_raw8_loss_model = build_raw8_combined_loss_model(model_by_id)
    combined_raw8_chart_path = CHART_DIR / "combined_tinygesture_finetune_jesternet_qat.png"
    if combined_raw8_loss_model:
        plot_loss_curve(combined_raw8_loss_model, combined_raw8_chart_path)
    scatter_path = CHART_DIR / "accuracy_vs_mmac.png"
    plot_accuracy_mmac(models, scatter_path, board_by_model)
    params_scatter_path = CHART_DIR / "accuracy_vs_params.png"
    plot_accuracy_params(models, params_scatter_path, board_by_model)
    params_stm32u5_scatter_path = CHART_DIR / "accuracy_vs_params_stm32u5.png"
    plot_accuracy_params(models, params_stm32u5_scatter_path, board_by_model, STM32U5_DEPLOYABLE_POINTS)

    doc = Document()
    configure_styles(doc)
    add_title(doc)

    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    p.add_run(" | Source: remote training logs under ")
    p.add_run(data["project"]).italic = True

    doc.add_heading("Executive Summary", level=1)
    add_bullet(doc, "The strongest quantized accuracy result in this report remains the huge-teacher QAT sweep at 77.555% Top-1 / 96.287% Top-5, using the signed-motion residual extra-conv student; that topology is still marked synth failed in the available placement attempt.")
    add_bullet(doc, "Signed motion consistently improved the 12-frame family: the residual model increased from 73.713% to 76.669% Top-1 after switching from absolute motion to signed motion.")
    add_bullet(doc, "Wider residual models improved accuracy at a moderate compute cost: the residual extra-conv student is about 56.79 MMAC per inference, compared with 30.84 MMAC for the simpler 12-plane motion baseline.")
    add_bullet(doc, "Teacher models reached higher software-only validation accuracy, but QAT compressed the deployable students back to the 76.7-77.6% Top-1 range.")
    add_bullet(doc, "A smaller signed residual distillation run finished on 2026-05-31: the student reached 77.757% Top-1 before QAT and 76.682% Top-1 / 95.942% Top-5 after QAT, with 393k parameters, 54.43 MMAC, and the previously measured 4.701 ms MAX78000 residual topology.")
    add_bullet(doc, "Board CNN time was measured on MAX78000FTHR known-answer firmware: JesterNet 8-frame is 2.760 ms, JesterMotion8/simple 12-channel is 3.754 ms, and the residual 12-channel topology is 4.701 ms.")
    add_bullet(doc, "The current true 3D CNN experiment is intentionally excluded from this report, per request.")

    doc.add_heading("Method and Quantization Pipeline", level=1)
    doc.add_paragraph(
        "The report consolidates epoch-level train loss and validation loss from PyTorch history files and ai8x training logs. "
        "For ai8x experiments, floating-point pretraining and quantization-aware training are treated as separate phases when both are present."
    )
    doc.add_paragraph(
        "Quantization flow: a floating-point or distilled checkpoint is first trained with ai8x-compatible layers. "
        "QAT then enables 8-bit weight/activation behavior, output shifts, clipping, and MAX78000-style rounding during training. "
        "The final QAT checkpoint is the source for fixed-point quantization and later ai8x synthesis. "
        "QAT success does not by itself guarantee placement on the MAX78000; synthesis/resource placement still needs to pass separately."
    )
    doc.add_paragraph(
        "MMAC definition: one multiply-accumulate is counted for Conv/Linear arithmetic in a single 64x64 inference. "
        "Pooling, activation, element-wise add, UART/camera preprocessing, and accelerator data movement are not included."
    )
    doc.add_paragraph(
        "CNN inference time in the summary table is the MAX78000 CNN_INFERENCE_TIMER value from known-answer firmware, "
        "built for BOARD=FTHR_RevA and repeated three times per measured topology. "
        "Rows sharing the same layer topology reuse the corresponding measured topology time. "
        "Rows marked not measured or synth fail do not have a board timing number in this report."
    )
    if board_measurements:
        add_table(
            doc,
            ["Measured topology", "Applies to", "MMAC", "Runs (us)", "Median ms"],
            [
                [
                    record.get("arch_group", record.get("model_id", "")),
                    ", ".join(record.get("applies_to") or []),
                    fmt_num(record.get("mmac"), 2),
                    ", ".join(str(v) for v in record.get("measured_us") or []),
                    fmt_num(record.get("median_ms"), 3),
                ]
                for record in board_measurements
            ],
            [0.95, 3.15, 0.55, 0.82, 0.55],
            font_size=6.2,
        )

    doc.add_heading("Final Accuracy and Resource Summary", level=1)
    doc.add_paragraph("Best Overall Top-1 may refer to a software teacher or pre-QAT student. QAT Top-1/Top-5 is the deployable quantized training phase when available.")
    add_table(
        doc,
        ["Model", "ai8x", "Params", "MMAC", "CNN ms", "Best Top-1", "Best Top-5", "QAT Top-1", "QAT Top-5"],
        make_accuracy_rows(models, board_by_model),
        [1.62, 0.34, 0.55, 0.44, 0.52, 0.62, 0.62, 0.62, 0.62],
        font_size=6.2,
    )

    doc.add_paragraph()
    doc.add_picture(str(scatter_path), width=Inches(6.25))
    cap = doc.add_paragraph(
        "Figure 1. Quantized Top-1 accuracy versus Conv/Linear MMAC for ai8x models. "
        "Color encodes deployment/synthesis status, marker shape encodes input/model category, and M/S/D tags mark motion, signed motion, and distillation."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_picture(str(params_scatter_path), width=Inches(6.25))
    cap = doc.add_paragraph(
        "Figure 2. Top-1 accuracy versus model parameters. The x-axis is logarithmic so the external "
        "TSM and saetlan/20BN-jester references can be shown alongside the sub-0.5M-parameter MAX78000 candidates. "
        "Color encodes deployment/synthesis status, marker shape encodes input/model category, and M/S/D tags mark motion, signed motion, and distillation."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "External reference note: saetlan/20BN-jester reports 85.99% top-1 for a 3D ResNet101 with dropout 0.5. "
        "Its parameter count is calculated from that repository's 3D ResNet101 bottleneck implementation and config "
        "as approximately 85.44M parameters. Lin et al. (ICCV 2019) report 97.0% top-1 on Jester for TSM with a ResNet-50 backbone; "
        "the paper reports 24.3M parameters for the 8-frame TSM ResNet-50 configuration and describes TSM as adding zero parameters. "
        "Both are included only as software reference points, not as deployable MAX78000 candidates."
    )
    doc.add_paragraph()
    doc.add_picture(str(params_stm32u5_scatter_path), width=Inches(6.25))
    cap = doc.add_paragraph(
        "Figure 2b. Top-1 accuracy versus model parameters with user-provided STM32 U5 deployable reference points. "
        "Purple triangles mark STM32 U5 models; all other encodings follow Figure 2."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "STM32 U5 note: the purple reference points are user-provided deployable model results: "
        "80K parameters at 61.4% Top-1, 273K parameters at 72.9% Top-1, and 722K parameters at 84.38% Top-1. "
        "They are included for cross-MCU context and are not MAX78000 measurements."
    )

    doc.add_heading("Validation Confusion Matrices", level=1)
    doc.add_paragraph(
        "The following normalized confusion matrices were regenerated from full validation-set ai8x --evaluate --confusion runs. "
        "Software-only teacher rows are not MAX78000/QAT deployment checkpoints and are not included in this figure set."
    )
    if confusion_summary:
        summary_rows = []
        for m in sorted(models, key=model_sort_key):
            row = confusion_summary.get(m["id"])
            if not row:
                continue
            summary_rows.append(
                [
                    m["display"],
                    row.get("samples", ""),
                    fmt_num(float(row["top1_percent"]), 3) if row.get("top1_percent") else "n/a",
                    fmt_num(float(row["mean_class_accuracy_percent"]), 3)
                    if row.get("mean_class_accuracy_percent")
                    else "n/a",
                ]
            )
        if summary_rows:
            add_table(
                doc,
                ["Model", "Samples", "Top-1", "Mean class acc."],
                summary_rows,
                [3.25, 0.7, 0.65, 0.8],
                font_size=6.4,
            )
    figure_index = 3
    for model in sorted(models, key=model_sort_key):
        if not model.get("confusion_path"):
            continue
        doc.add_heading(model["display"], level=2)
        doc.add_picture(model["confusion_path"], width=Inches(6.25))
        cap = doc.add_paragraph(f"Figure {figure_index}. Normalized validation confusion matrix for {model['display']}.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure_index += 1
        doc.add_page_break()
    training_figure_base = figure_index

    doc.add_heading("Training and QAT Time Summary", level=1)
    doc.add_paragraph(
        "Timing is wall-clock when wrapper timestamps were available. For copied QAT logs without an end timestamp, duration is estimated from logged per-batch training time; validation and checkpoint overhead are therefore undercounted slightly."
    )
    add_table(
        doc,
        ["Model", "Phase timing", "Total"],
        make_time_rows(models),
        [1.95, 3.62, 0.68],
        font_size=6.5,
    )

    doc.add_heading("Training Loss Curves", level=1)
    doc.add_paragraph(
        "Each chart shows train loss and validation loss versus epoch. When a run has multiple phases, the x-axis is concatenated by phase and vertical separators indicate phase boundaries."
    )
    figure_no = training_figure_base
    curves_on_page = 0
    skip_loss_ids: set[str] = set()
    for model in models:
        if model["id"] in skip_loss_ids:
            continue
        if model["id"] == COMBINED_RAW8_LOSS_IDS[0] and combined_raw8_loss_model:
            doc.add_heading("TinyGesture raw 8-frame finetune + JesterNet raw 8-frame ai8x QAT", level=2)
            desc = (
                "Combined view of the raw 8-frame software finetune and the following ai8x QAT run. "
                f"The finetune portion is cropped to the first {LOSS_PLOT_EPOCH_LIMITS[COMBINED_RAW8_LOSS_IDS[0]]} epochs; "
                "the QAT portion is shown in full."
            )
            add_training_curve_figure(
                doc,
                str(combined_raw8_chart_path),
                "Train and validation loss for TinyGesture raw 8-frame 300-epoch finetune and JesterNet raw 8-frame ai8x QAT. Finetune shows the first 80 epochs.",
                figure_no,
                desc,
            )
            figure_no += 1
            curves_on_page += 1
            skip_loss_ids.add(COMBINED_RAW8_LOSS_IDS[1])
            if curves_on_page % 2 == 0:
                doc.add_page_break()
            continue
        doc.add_heading(model["display"], level=2)
        desc = (
            f"Architecture: {model['arch']}. Input: {model['input']}. "
            f"Params: {fmt_num(model.get('params'), 0)}. MMAC: {fmt_num(model.get('mmac'), 2)}. "
            f"Best overall Top-1: {fmt_pct((model.get('best') or {}).get('top1'))}. "
        )
        q = best_qat(model)
        if q:
            desc += f"Best QAT Top-1/Top-5: {fmt_pct(q.get('top1'))} / {fmt_pct(q.get('top5'))}."
        if model.get("notes"):
            desc += f" Note: {model['notes']}"
        epoch_limit = LOSS_PLOT_EPOCH_LIMITS.get(model.get("id", ""))
        limit_note = f" First {epoch_limit} epochs shown." if epoch_limit else ""
        add_training_curve_figure(
            doc,
            model["chart_path"],
            f"Train and validation loss for {model['display']}.{limit_note}",
            figure_no,
            desc,
        )
        figure_no += 1
        curves_on_page += 1
        if curves_on_page % 2 == 0 and model != models[-1]:
            doc.add_page_break()

    doc.add_heading("Limitations", level=1)
    add_bullet(doc, "The report is log-derived. If a run did not record wall-clock timestamps, the training-time field is marked as not logged or estimated from batch timing.")
    add_bullet(doc, "MMAC is an arithmetic proxy, not an energy number. The CNN ms column is real board timing only where a measured known-answer topology exists.")
    add_bullet(doc, "Teacher rows are software-only and are included to explain distillation; they are not MAX78000 deployment candidates.")
    add_bullet(doc, "Board placement status is separate from QAT accuracy. The residual extra-conv family hit kernel-memory placement failure in the available synthesis attempt, so those rows are accuracy results only until placement is resolved.")

    doc.save(DOCX_PATH)
    print(DOCX_PATH.resolve())


if __name__ == "__main__":
    main()
